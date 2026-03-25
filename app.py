import os
import json
import io
import base64
import tempfile
from flask import Flask, render_template, request, jsonify, session, send_file
from groq import Groq
from PIL import Image
from docx import Document
from twilio.rest import Client

# Local Imports
from translations import translations
from products import products

# Suppress TF ops if any (though we removed TF, some envs might still have it)
os.environ['TF_ENABLE_ONEDNN_OPTS'] = '0'
os.environ['TF_CPP_MIN_LOG_LEVEL'] = '2'

app = Flask(__name__)
# Secure secret key
app.secret_key = os.environ.get("FLASK_SECRET_KEY", "agrisolution_secret_123_prod")

# Load Config
WORKING_DIR = os.path.dirname(os.path.abspath(__file__))
config_data = {}
try:
    config_path = os.path.join(WORKING_DIR, "config.json")
    if os.path.exists(config_path):
        with open(config_path) as f:
            config_data = json.load(f)
except Exception as e:
    print(f"Error loading config: {e}")

GROQ_API_KEY = config_data.get("GROQ_API_KEY", "")
TWILIO_ACCOUNT_SID = config_data.get("TWILIO_ACCOUNT_SID", "")
TWILIO_AUTH_TOKEN = config_data.get("TWILIO_AUTH_TOKEN", "")
TWILIO_PHONE_NUMBER = config_data.get("TWILIO_PHONE_NUMBER", "")
YOUR_PHONE_NUMBER = config_data.get("YOUR_PHONE_NUMBER", "")

# Groq Vision Configuration
# meta-llama/llama-4-scout-17b-16e-instruct is currently the recommended vision-capable model on GroqCloud
GROQ_VISION_MODEL = "meta-llama/llama-4-scout-17b-16e-instruct"

def encode_image(image_file):
    """Encode PIL image to base64 string directly from memory."""
    buffered = io.BytesIO()
    image_file.save(buffered, format="JPEG")
    return base64.b64encode(buffered.getvalue()).decode('utf-8')

@app.before_request
def ensure_language():
    """Ensure a default language is set in the session."""
    if 'lang' not in session:
        session['lang'] = 'en'

@app.route('/')
def index():
    t = translations[session['lang']]
    return render_template('index.html', t=t, lang=session['lang'])

@app.route('/diagnose')
def diagnose():
    t = translations[session['lang']]
    return render_template('diagnose.html', t=t, lang=session['lang'])

@app.route('/planty')
def planty():
    t = translations[session['lang']]
    return render_template('planty.html', t=t, lang=session['lang'])

@app.route('/shop')
def shop_page():
    lang = session.get('lang', 'en')
    t = translations[lang]
    return render_template('shop.html', t=t, lang=lang, products=products)

@app.route('/contact')
def contact_page():
    t = translations[session['lang']]
    return render_template('contact.html', t=t, lang=session['lang'])

@app.route('/set_language', methods=['POST'])
def set_language():
    lang = request.json.get('lang', 'en')
    if lang in translations:
        session['lang'] = lang
        return jsonify({"status": "success"})
    return jsonify({"status": "error", "message": "Invalid language"}), 400

@app.route('/predict', methods=['POST'])
def predict():
    if 'file' not in request.files:
        return jsonify({"error": "No file uploaded"}), 400
    
    file = request.files['file']
    try:
        # Open image to ensure it's valid
        image = Image.open(file.stream)
        
        # Reset file stream for encoding
        file.stream.seek(0)
        base64_image = encode_image(image)
        
        t = translations[session['lang']]
        
        # Groq Vision API Call
        client = Groq(api_key=GROQ_API_KEY)
        
        prompt = (
            f"Analyze the attached image as an agricultural expert. \n"
            f"1. If the image is NOT a plant leaf or is a 'wrong image', respond exactly with: 'Disease: Incorrect Image\nTreatment: Incorrect image, please upload a valid plant leaf.'\n"
            f"2. If it is a plant, identify the disease (or 'Healthy') and give 2-3 lines of treatment advice in {session['lang']}.\n"
            f"Format:\nDisease: [Name]\nTreatment: [Advice]"
        )

        response = client.chat.completions.create(
            model=GROQ_VISION_MODEL,
            messages=[
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": prompt},
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:image/jpeg;base64,{base64_image}",
                            },
                        },
                    ],
                }
            ],
            max_tokens=250,
        )

        result_text = response.choices[0].message.content.strip()
        
        # Parse result
        display_name = "Detected Condition"
        solution_info = result_text
        
        if "Disease:" in result_text and "Treatment:" in result_text:
            parts = result_text.split("Treatment:")
            display_name = parts[0].replace("Disease:", "").strip()
            solution_info = parts[1].strip()

        # Handle validation failure
        status = "success"
        if "Incorrect Image" in display_name:
            status = "invalid_image"

        # Store last prediction in session for report generation
        session['last_prediction'] = {
            'class': display_name,
            'confidence': "AI Vision Analysis",
            'solution': solution_info,
            'display_name': display_name
        }

        return jsonify({
            "predicted_class": display_name,
            "display_name": display_name,
            "confidence": "Vision Analysis",
            "solution": solution_info,
            "status": status
        })
    except Exception as e:
        print(f"Vision API Error: {str(e)}")
        # If API fails, provide a fallback message
        return jsonify({"error": f"Analysis unavailable: {str(e)}"}), 500

@app.route('/chat', methods=['POST'])
def chat():
    user_message = request.json.get('message', '')
    if not user_message:
        return jsonify({"error": "Empty message"}), 400
    
    chat_history = session.get('chat_history', [])
    chat_history.append({"role": "user", "content": user_message})
    
    try:
        client = Groq(api_key=GROQ_API_KEY)
        messages = [
            {"role": "system", "content": "You are an expert in tomato diseases and treatment solutions"},
            *chat_history
        ]
        response = client.chat.completions.create(
            model="llama-3.1-8b-instant",
            messages=messages
        )
        assistant_response = response.choices[0].message.content
        chat_history.append({"role": "assistant", "content": assistant_response})
        session['chat_history'] = chat_history
        return jsonify({"response": assistant_response})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/clear_chat', methods=['POST'])
def clear_chat():
    session['chat_history'] = []
    return jsonify({"status": "success"})

@app.route('/generate_report')
def generate_report():
    last_pred = session.get('last_prediction')
    if not last_pred:
        return "No prediction data found", 400
    
    t = translations[session['lang']]
    doc = Document()
    doc.add_heading(t["title_doc"], level=1)
    doc.add_paragraph(f"**{t['disease_detected']}** {last_pred['display_name']}")
    doc.add_paragraph(f"**{t['confidence']}** {last_pred['confidence']}")
    doc.add_heading(t["solution_info"], level=2)
    doc.add_paragraph(last_pred['solution'])
    
    temp_doc = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp_doc.name)
    temp_doc.close()
    
    return send_file(temp_doc.name, as_attachment=True, download_name="Diagnosis_Report.docx")

@app.route('/send_contact', methods=['POST'])
def send_contact():
    data = request.form
    name = data.get('name')
    email = data.get('email')
    phone = data.get('phone')
    message = data.get('message')
    
    if not all([name, email, message]):
        return jsonify({"status": "error", "message": "Missing required fields"}), 400

    try:
        client = Client(TWILIO_ACCOUNT_SID, TWILIO_AUTH_TOKEN)
        sms_body = (
            f"🚜 New Message from Farmer:\n"
            f"👤 Name: {name}\n"
            f"📧 Email: {email}\n"
            f"📞 Phone: {phone}\n"
            f"🌾 Message: {message}"
        )
        client.messages.create(
            body=sms_body,
            from_=TWILIO_PHONE_NUMBER,
            to=YOUR_PHONE_NUMBER
        )
        return jsonify({"status": "success"})
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500

if __name__ == '__main__':
    app.run(debug=True, port=5000)
